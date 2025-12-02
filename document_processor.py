# document_processor.py

import os
from pathlib import Path
from typing import List, Dict, Any
from langchain_core.documents import Document
from langchain_community.document_loaders import UnstructuredPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
import pandas as pd
from docx import Document as DocxDocument
import config


class DocumentProcessor:
    """
    Unified document processor for PDF, Excel, and Word files.
    Extracts content with rich metadata for source citation.
    """
    
    def __init__(self, chunk_size: int = None, chunk_overlap: int = None):
        """
        Initialize the document processor.
        
        Args:
            chunk_size: Size of text chunks (default from config)
            chunk_overlap: Overlap between chunks (default from config)
        """
        self.chunk_size = chunk_size or config.CHUNK_SIZE
        self.chunk_overlap = chunk_overlap or config.CHUNK_OVERLAP
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def load_document(self, file_path: str) -> List[Document]:
        """
        Load a document and return chunks with metadata.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of Document objects with content and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self._load_pdf(file_path)
        elif extension in ['.xlsx', '.xls']:
            return self._load_excel(file_path)
        elif extension in ['.docx', '.doc']:
            return self._load_word(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    def _load_pdf(self, file_path: Path) -> List[Document]:
        """Load and process PDF file."""
        try:
            loader = UnstructuredPDFLoader(str(file_path))
            pages = loader.load()
            
            documents = []
            doc_type = config.get_document_type(file_path.name)
            
            for page in pages:
                # Extract page number from metadata if available
                page_num = page.metadata.get('page_number', 'Unknown')
                
                # Create base metadata
                base_metadata = {
                    'source': file_path.name,
                    'source_type': 'PDF',
                    'document_type': doc_type,
                    'page': page_num,
                    'file_path': str(file_path)
                }
                
                # Split page content into chunks
                chunks = self.text_splitter.split_text(page.page_content)
                
                for i, chunk in enumerate(chunks):
                    chunk_metadata = base_metadata.copy()
                    chunk_metadata['chunk_index'] = i
                    
                    documents.append(Document(
                        page_content=chunk,
                        metadata=chunk_metadata
                    ))
            
            return documents
            
        except Exception as e:
            raise Exception(f"Error loading PDF {file_path.name}: {str(e)}")
    
    def _load_excel(self, file_path: Path) -> List[Document]:
        """Load and process Excel file."""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            documents = []
            doc_type = config.get_document_type(file_path.name)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                # Convert DataFrame to text representation
                # Include column headers and data
                text_content = f"Feuille: {sheet_name}\n\n"
                
                # Add headers
                text_content += "Colonnes: " + ", ".join(df.columns.astype(str)) + "\n\n"
                
                # Add data row by row
                for idx, row in df.iterrows():
                    row_text = " | ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                    if row_text.strip():
                        text_content += f"Ligne {idx + 2}: {row_text}\n"
                
                # Create base metadata
                base_metadata = {
                    'source': file_path.name,
                    'source_type': 'Excel',
                    'document_type': doc_type,
                    'sheet': sheet_name,
                    'file_path': str(file_path)
                }
                
                # Split content into chunks
                chunks = self.text_splitter.split_text(text_content)
                
                for i, chunk in enumerate(chunks):
                    chunk_metadata = base_metadata.copy()
                    chunk_metadata['chunk_index'] = i
                    
                    documents.append(Document(
                        page_content=chunk,
                        metadata=chunk_metadata
                    ))
            
            return documents
            
        except Exception as e:
            raise Exception(f"Error loading Excel {file_path.name}: {str(e)}")
    
    def _load_word(self, file_path: Path) -> List[Document]:
        """Load and process Word document."""
        try:
            doc = DocxDocument(str(file_path))
            documents = []
            doc_type = config.get_document_type(file_path.name)
            
            # Extract text from paragraphs
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        full_text.append(row_text)
            
            content = "\n".join(full_text)
            
            # Create base metadata
            base_metadata = {
                'source': file_path.name,
                'source_type': 'Word',
                'document_type': doc_type,
                'file_path': str(file_path)
            }
            
            # Split content into chunks
            chunks = self.text_splitter.split_text(content)
            
            for i, chunk in enumerate(chunks):
                chunk_metadata = base_metadata.copy()
                chunk_metadata['chunk_index'] = i
                chunk_metadata['section'] = f"Partie {i + 1}"
                
                documents.append(Document(
                    page_content=chunk,
                    metadata=chunk_metadata
                ))
            
            return documents
            
        except Exception as e:
            raise Exception(f"Error loading Word document {file_path.name}: {str(e)}")
    
    def load_directory(self, directory_path: str) -> List[Document]:
        """
        Load all supported documents from a directory recursively.
        
        Args:
            directory_path: Path to the directory
            
        Returns:
            List of all Document objects from all files
        """
        directory = Path(directory_path)
        all_documents = []
        processed_files = []
        failed_files = []
        
        # Find all supported files recursively
        for ext in config.ALL_SUPPORTED_EXTENSIONS:
            for file_path in directory.rglob(f"*{ext}"):
                try:
                    print(f"📄 Processing: {file_path.name}")
                    docs = self.load_document(file_path)
                    all_documents.extend(docs)
                    processed_files.append(file_path.name)
                    print(f"   ✅ Loaded {len(docs)} chunks")
                except Exception as e:
                    print(f"   ❌ Failed: {str(e)}")
                    failed_files.append((file_path.name, str(e)))
        
        print(f"\n📊 Summary:")
        print(f"   ✅ Successfully processed: {len(processed_files)} files")
        print(f"   ❌ Failed: {len(failed_files)} files")
        print(f"   📝 Total chunks created: {len(all_documents)}")
        
        if failed_files:
            print(f"\n⚠️ Failed files:")
            for filename, error in failed_files:
                print(f"   - {filename}: {error}")
        
        return all_documents
